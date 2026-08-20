"""Export a transcription as .txt, .md, .docx, or .pdf.

All exporters return bytes so the FastAPI endpoint can stream them straight
back as a download.
"""
import io
import re

from docx import Document
from docx.shared import Pt
from fpdf import FPDF


def _safe_filename(title: str, ext: str) -> str:
    """Trim to a filesystem-friendly filename."""
    slug = re.sub(r"[^A-Za-z0-9 _.-]", "", title or "transcription").strip()
    slug = re.sub(r"\s+", "-", slug)[:60] or "transcription"
    return f"{slug}.{ext}"


def as_txt(text: str, title: str) -> tuple[bytes, str, str]:
    body = f"{title}\n{'=' * len(title)}\n\n{text}\n"
    return body.encode("utf-8"), _safe_filename(title, "txt"), "text/plain"


def as_md(text: str, title: str) -> tuple[bytes, str, str]:
    body = f"# {title}\n\n{text}\n"
    return body.encode("utf-8"), _safe_filename(title, "md"), "text/markdown"


def as_docx(text: str, title: str) -> tuple[bytes, str, str]:
    doc = Document()

    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.name = "Calibri"

    for paragraph in text.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        p = doc.add_paragraph()
        for line in paragraph.split("\n"):
            if p.runs:
                p.add_run("\n")
            run = p.add_run(line)
            run.font.name = "Calibri"
            run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return (
        buf.getvalue(),
        _safe_filename(title, "docx"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def as_pdf(text: str, title: str) -> tuple[bytes, str, str]:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    # fpdf2's core fonts are latin-1; strip anything it cannot render so the
    # library does not raise on curly quotes or emoji.
    def _safe(s: str) -> str:
        return s.encode("latin-1", "replace").decode("latin-1")

    pdf.set_font("Helvetica", "B", 18)
    pdf.multi_cell(0, 10, _safe(title))
    pdf.ln(4)

    pdf.set_font("Helvetica", "", 11)
    for paragraph in text.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        pdf.multi_cell(0, 6, _safe(paragraph))
        pdf.ln(2)

    out = bytes(pdf.output())
    return out, _safe_filename(title, "pdf"), "application/pdf"


EXPORTERS = {
    "txt": as_txt,
    "md": as_md,
    "docx": as_docx,
    "pdf": as_pdf,
}
